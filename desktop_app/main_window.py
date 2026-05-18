from __future__ import annotations

import json
import re
import shutil
import traceback
import webbrowser
from dataclasses import replace
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from PySide6.QtCore import QEvent, QModelIndex, QObject, QRect, QSize, QTimer, Qt, Slot
from PySide6.QtGui import QAction, QColor, QFont, QIcon, QKeySequence, QPainter, QPixmap
from PySide6.QtWidgets import (
    QAbstractItemView,
    QApplication,
    QButtonGroup,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFrame,
    QHBoxLayout,
    QHeaderView,
    QInputDialog,
    QLabel,
    QListWidget,
    QListWidgetItem,
    QMainWindow,
    QMenu,
    QMessageBox,
    QProgressBar,
    QPushButton,
    QScrollArea,
    QSizePolicy,
    QStatusBar,
    QTableView,
    QTableWidget,
    QTableWidgetItem,
    QTextEdit,
    QStyle,
    QStyledItemDelegate,
    QToolButton,
    QStyleOptionViewItem,
    QVBoxLayout,
    QWidget,
)

from etp_client import EtpClient, PROCEDURE_TYPE_OPTIONS, STATUS_OPTIONS, step_id_label, trend_pur_label
from gpb_business_client import GpbBusinessClient
from roseltorg_client import (
    ROSELTORG_PROCEDURE_TYPE_OPTIONS,
    ROSELTORG_SEARCH_BY_OPTIONS,
    ROSELTORG_STATUS_OPTIONS,
    RoseltorgClient,
)

from .assets import asset_path
from .constants import (
    ANALYSIS_DIR,
    APP_TITLE,
    CACHE_FILE,
    COLUMNS,
    DOCUMENTS_DIR,
    KEYWORDS_FILE,
    LM_STUDIO_BASE_URL,
    LM_STUDIO_MODEL,
    VIEW_URL,
    user_writable_root,
)
from .lm_table_analysis import ANALYSIS_TABLE_HEADERS_RU
from .gpb_rag.chat import answer_question_from_saved_index
from .keywords import load_keyword_items, parse_keyword_items, save_keyword_items
from .models import ProcedureFilterProxy, ProcedureTableModel
from .params import ClientFilters, SearchParams
from .sidebar import Sidebar
from .utils import fmt_date, parse_dt, parse_price
from .worker import (
    TaskRunner,
    make_analyze_procedure_task,
    make_download_documents_task,
    make_search_task,
)


DELETED_TENDERS_FILE = user_writable_root() / "deleted_tenders.json"


class LimitedWrapDelegate(QStyledItemDelegate):
    """Переносит длинный текст в таблице, но не даёт строкам занять весь экран."""

    MAX_ROW_HEIGHT = 74
    MIN_ROW_HEIGHT = 44
    PADDING = 14

    BADGE_COLUMNS = {1, 5, 10}

    def initStyleOption(self, option: QStyleOptionViewItem, index: QModelIndex) -> None:  # noqa: N802
        super().initStyleOption(option, index)
        option.features |= QStyleOptionViewItem.ViewItemFeature.WrapText
        option.textElideMode = Qt.TextElideMode.ElideRight

    def sizeHint(self, option: QStyleOptionViewItem, index: QModelIndex) -> QSize:  # noqa: N802
        base = super().sizeHint(option, index)
        text = str(index.data(Qt.DisplayRole) or "")
        if not text:
            return QSize(base.width(), self.MIN_ROW_HEIGHT)

        width = option.rect.width()
        view = self.parent()
        if width <= 0 and isinstance(view, QTableView):
            width = view.columnWidth(index.column())
        text_width = max(40, width - self.PADDING)
        rect = option.fontMetrics.boundingRect(
            0,
            0,
            text_width,
            10_000,
            int(Qt.TextFlag.TextWordWrap | Qt.TextFlag.TextExpandTabs),
            text,
        )
        height = min(max(self.MIN_ROW_HEIGHT, rect.height() + self.PADDING), self.MAX_ROW_HEIGHT)
        return QSize(base.width(), height)

    def _badge_colors(self, column: int, text: str) -> tuple[QColor, QColor, QColor]:
        t = text.casefold()
        if column == 5:
            if "223" in t:
                return QColor("#0f5132"), QColor("#1f9d55"), QColor("#bdf7d3")
            return QColor("#123a74"), QColor("#2563eb"), QColor("#c9ddff")
        if column == 10:
            if "отмен" in t:
                return QColor("#4a1625"), QColor("#9f2944"), QColor("#ffc2cf")
            if "итог" in t or "рассмотр" in t or "подвед" in t:
                return QColor("#4a3511"), QColor("#9a6a12"), QColor("#ffe1a3")
            return QColor("#123a74"), QColor("#2563eb"), QColor("#c9ddff")
        if "аукцион" in t:
            return QColor("#073744"), QColor("#0891b2"), QColor("#bff4ff")
        if "запрос" in t:
            return QColor("#102f69"), QColor("#2563eb"), QColor("#d6e4ff")
        if "открыт" in t:
            return QColor("#4a3208"), QColor("#b7791f"), QColor("#ffe4a6")
        return QColor("#211d55"), QColor("#4f46e5"), QColor("#e1ddff")

    def paint(self, painter: QPainter, option: QStyleOptionViewItem, index: QModelIndex) -> None:  # noqa: N802
        text = str(index.data(Qt.DisplayRole) or "")
        if index.column() not in self.BADGE_COLUMNS or not text.strip():
            super().paint(painter, option, index)
            return

        painter.save()
        selected = bool(option.state & QStyle.StateFlag.State_Selected)
        base = QColor("#263168" if selected else ("#0a1128" if index.row() % 2 else "#070d20"))
        painter.fillRect(option.rect, base)

        bg, border, fg = self._badge_colors(index.column(), text)
        metrics = option.fontMetrics
        badge_text = metrics.elidedText(text, Qt.TextElideMode.ElideRight, max(24, option.rect.width() - 18))
        text_w = min(metrics.horizontalAdvance(badge_text) + 18, option.rect.width() - 12)
        badge_h = min(24, option.rect.height() - 12)
        badge = QRect(
            option.rect.x() + 8,
            option.rect.y() + max(6, (option.rect.height() - badge_h) // 2),
            max(24, text_w),
            badge_h,
        )
        painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)
        painter.setPen(border)
        painter.setBrush(bg)
        painter.drawRoundedRect(badge, 7, 7)
        painter.setPen(fg)
        painter.drawText(badge.adjusted(9, 0, -8, 0), Qt.AlignmentFlag.AlignVCenter | Qt.AlignmentFlag.AlignLeft, badge_text)
        painter.restore()


class MainWindow(QMainWindow):
    def __init__(self) -> None:
        super().__init__()
        self.setWindowTitle(APP_TITLE)
        self.setWindowIcon(QIcon(str(asset_path("iconfounder.png"))))
        self.resize(1500, 900)

        self.client = EtpClient()
        self.runner = TaskRunner(self)
        self.model = ProcedureTableModel(self)
        self.proxy = ProcedureFilterProxy(self)
        self.proxy.setSourceModel(self.model)

        self._last_total: int = 0
        self._current_start: int = 0
        self._last_user: Optional[str] = None
        self._cache_dirty: bool = False
        self._cache_save_enabled: bool = True
        self._platform_key: str = "gpb"
        self._api_debug_chunks: list[str] = []
        self._deleted_tender_records: dict[str, dict[str, Any]] = self._load_deleted_tenders()
        self._deleted_tender_keys: set[str] = set(self._deleted_tender_records)

        self._cache_save_timer = QTimer(self)
        self._cache_save_timer.setSingleShot(True)
        self._cache_save_timer.timeout.connect(self._save_cache_now)
        self._analysis_sink: dict[str, Any] = {}

        self._build_ui()
        self._announce_cache_on_start()
        self._update_controls()

    # ------------------------------------------------------------------ UI
    def _build_ui(self) -> None:
        # ---------- Верхняя панель ----------
        top = QFrame()
        top.setObjectName("TopBar")
        top.setFixedHeight(78)
        top_layout = QHBoxLayout(top)
        top_layout.setContentsMargins(22, 10, 22, 10)
        top_layout.setSpacing(18)

        platform_switcher = QFrame()
        platform_switcher.setObjectName("PlatformSwitcher")
        platform_switcher.setFixedHeight(48)
        platform_layout = QHBoxLayout(platform_switcher)
        platform_layout.setContentsMargins(4, 4, 4, 4)
        platform_layout.setSpacing(4)
        self.platform_group = QButtonGroup(self)
        self.platform_group.setExclusive(True)
        self.btn_platform_gpb = QPushButton("ГПБ")
        self.btn_platform_gpb.setObjectName("PlatformButton")
        self.btn_platform_gpb.setCheckable(True)
        self.btn_platform_gpb.setChecked(True)
        self.btn_platform_gpb.setFixedHeight(40)
        self.gpb_platform_menu = QMenu(self.btn_platform_gpb)
        self.act_platform_etp_gpb = self.gpb_platform_menu.addAction("ЭТП ГПБ")
        self.act_platform_gpb_business = self.gpb_platform_menu.addAction("ГПБ Бизнес")
        self.btn_platform_gpb.setMenu(self.gpb_platform_menu)
        self.btn_platform_roseltorg = QPushButton("Росэлторг")
        self.btn_platform_roseltorg.setObjectName("PlatformButton")
        self.btn_platform_roseltorg.setCheckable(True)
        self.btn_platform_roseltorg.setFixedHeight(40)
        self.platform_group.addButton(self.btn_platform_gpb)
        self.platform_group.addButton(self.btn_platform_roseltorg)
        platform_layout.addWidget(self.btn_platform_gpb)
        platform_layout.addWidget(self.btn_platform_roseltorg)
        self.act_platform_etp_gpb.triggered.connect(lambda: self._select_platform("gpb"))
        self.act_platform_gpb_business.triggered.connect(lambda: self._select_platform("gpb_business"))
        self.btn_platform_roseltorg.clicked.connect(lambda: self._select_platform("roseltorg"))
        top_layout.addWidget(platform_switcher, 0, Qt.AlignVCenter)

        t_title_box = QVBoxLayout()
        t_title_box.setSpacing(0)
        self.title_label = QLabel("ЭТП ГПБ — Актуальные процедуры")
        self.title_label.setObjectName("TopBarTitle")
        self.subtitle_label = QLabel("Поиск, фильтры и экспорт")
        self.subtitle_label.setObjectName("TopBarSubtitle")
        t_title_box.addWidget(self.title_label)
        t_title_box.addWidget(self.subtitle_label)
        top_layout.addLayout(t_title_box)
        top_layout.addStretch(1)

        self.user_label = QLabel("Пользователь: —")
        self.user_label.setStyleSheet("color: #9ca8d7;")
        top_layout.addWidget(self.user_label)

        self.session_badge = QLabel("○  Браузер не запущен")
        self.session_badge.setObjectName("SessionBadge")
        self.session_badge.setProperty("ok", "idle")
        top_layout.addWidget(self.session_badge)

        # ---------- Фильтры и основная область ----------
        self.sidebar = Sidebar()
        self.sidebar.searchRequested.connect(self._on_search)
        self.sidebar.resetRequested.connect(self._on_reset_filters)
        self.sidebar.clientFiltersChanged.connect(self._on_filters_changed)
        self.sidebar.editKeywordsRequested.connect(self._on_edit_keywords)

        main_area = QWidget()
        main_area.setMinimumHeight(360)
        main_area_layout = QVBoxLayout(main_area)
        main_area_layout.setContentsMargins(14, 12, 14, 10)
        main_area_layout.setSpacing(10)

        self.cache_banner = QFrame()
        self.cache_banner.setObjectName("CacheBanner")
        cache_layout = QHBoxLayout(self.cache_banner)
        cache_layout.setContentsMargins(16, 10, 10, 10)
        cache_layout.setSpacing(10)
        self.cache_banner_icon = QLabel("ⓘ")
        self.cache_banner_icon.setObjectName("CacheBannerIcon")
        cache_layout.addWidget(self.cache_banner_icon)
        self.lbl_counter = QLabel("Данных нет. Нажмите «Поиск».")
        self.lbl_counter.setObjectName("CacheBannerText")
        cache_layout.addWidget(self.lbl_counter, 1)
        self.btn_cache_dismiss = QToolButton()
        self.btn_cache_dismiss.setObjectName("CacheBannerClose")
        self.btn_cache_dismiss.setText("×")
        self.btn_cache_dismiss.setCursor(Qt.PointingHandCursor)
        self.btn_cache_dismiss.clicked.connect(lambda: self.cache_banner.setVisible(False))
        cache_layout.addWidget(self.btn_cache_dismiss)
        self.cache_banner.setVisible(False)
        main_area_layout.addWidget(self.cache_banner)

        actions = QFrame()
        actions.setObjectName("ActionsBar")
        actions_layout = QHBoxLayout(actions)
        actions_layout.setContentsMargins(2, 4, 2, 4)
        actions_layout.setSpacing(8)
        actions_layout.addStretch(1)

        self.btn_export = QPushButton("Экспорт в XLSX")
        self.btn_export.setObjectName("Ghost")
        self.btn_export.setIcon(QIcon(str(asset_path("xls.png"))))
        self.btn_export.setIconSize(QSize(16, 16))
        self.btn_export.clicked.connect(self._on_export)
        actions_layout.addWidget(self.btn_export)

        self.btn_save_api_debug = QPushButton("Сохранить API-логи")
        self.btn_save_api_debug.setObjectName("Ghost")
        self.btn_save_api_debug.setIcon(QIcon(str(asset_path("log.png"))))
        self.btn_save_api_debug.setIconSize(QSize(16, 16))
        self.btn_save_api_debug.setToolTip("Сохранить запросы, headers, body, token и ответы API в файл")
        self.btn_save_api_debug.clicked.connect(self._save_api_debug)
        self.btn_save_api_debug.setEnabled(False)
        actions_layout.addWidget(self.btn_save_api_debug)

        main_area_layout.addWidget(actions)

        # Таблица
        self.table = QTableView()
        self.table.setModel(self.proxy)
        self.table.setSortingEnabled(True)
        self.table.sortByColumn(8, Qt.AscendingOrder)
        self.table.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.table.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.table.setHorizontalScrollMode(QAbstractItemView.ScrollMode.ScrollPerPixel)
        self.table.setVerticalScrollMode(QAbstractItemView.ScrollMode.ScrollPerPixel)
        # Иначе layout раздувает окно по широкому sizeHint таблицы; скролл — только внутри QTableView.
        self.table.setSizePolicy(
            QSizePolicy.Policy.Ignored,
            QSizePolicy.Policy.Expanding,
        )
        self.table.setSelectionBehavior(QTableView.SelectRows)
        self.table.setSelectionMode(QTableView.ExtendedSelection)
        self.table.setAlternatingRowColors(True)
        self.table.verticalHeader().setDefaultSectionSize(44)
        self.table.verticalHeader().setFixedWidth(30)
        self.table.verticalHeader().setHighlightSections(False)
        hh = self.table.horizontalHeader()
        hh.setMinimumHeight(44)
        hh.setStretchLastSection(False)
        hh.setCascadingSectionResizes(False)
        self.table.setWordWrap(True)
        self.table.setItemDelegate(LimitedWrapDelegate(self.table))
        hh.sectionResized.connect(lambda *_: self._schedule_table_row_resize())
        self.table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self._on_context_menu)
        self.table.doubleClicked.connect(self._on_row_double_clicked)
        self.proxy.modelReset.connect(self._apply_table_column_widths)
        self.proxy.modelReset.connect(self._schedule_table_row_resize)
        self.proxy.rowsInserted.connect(self._schedule_table_row_resize)
        self.proxy.layoutChanged.connect(self._schedule_table_row_resize)
        self.table.viewport().installEventFilter(self)
        self._apply_table_column_widths()
        self._build_empty_state()
        main_area_layout.addWidget(self.table, 1)

        bottom_bar = QFrame()
        bottom_bar.setObjectName("BottomBar")
        bottom_layout = QHBoxLayout(bottom_bar)
        bottom_layout.setContentsMargins(10, 8, 10, 8)
        bottom_layout.setSpacing(8)

        self.btn_prev_page = QPushButton("←")
        self.btn_prev_page.setToolTip("Предыдущая страница")
        self.btn_prev_page.clicked.connect(self._on_prev_page)
        self.btn_prev_page.setEnabled(False)
        bottom_layout.addWidget(self.btn_prev_page)

        self.lbl_page = QLabel("Страница 0 из 0")
        self.lbl_page.setMinimumWidth(120)
        self.lbl_page.setAlignment(Qt.AlignCenter)
        bottom_layout.addWidget(self.lbl_page)

        self.btn_next_page = QPushButton("→")
        self.btn_next_page.setToolTip("Следующая страница")
        self.btn_next_page.clicked.connect(self._on_next_page)
        self.btn_next_page.setEnabled(False)
        bottom_layout.addWidget(self.btn_next_page)

        self.btn_download_docs = QPushButton("Скачать документы")
        self.btn_download_docs.setToolTip("Скачать документацию выбранных процедур")
        self.btn_download_docs.clicked.connect(self._on_download_documents)
        self.btn_download_docs.setEnabled(False)
        bottom_layout.addWidget(self.btn_download_docs)

        self.btn_analyze = QPushButton("Проанализировать")
        self.btn_analyze.setToolTip(
            "Собрать текст карточки с ЭТП ГПБ и отправить в LM Studio для заполнения таблицы анализа"
        )
        self.btn_analyze.clicked.connect(self._on_analyze_procedures)
        self.btn_analyze.setEnabled(False)
        bottom_layout.addWidget(self.btn_analyze)

        self.btn_show_analysis = QPushButton("Результат анализа")
        self.btn_show_analysis.setToolTip("Повторно открыть последнее окно результата анализа")
        self.btn_show_analysis.clicked.connect(self._on_show_analysis_result)
        self.btn_show_analysis.setEnabled(False)
        bottom_layout.addWidget(self.btn_show_analysis)

        bottom_layout.addStretch(1)

        self.btn_blacklist = QPushButton("Черный список")
        self.btn_blacklist.setToolTip("Показать удалённые тендеры и восстановить выбранные")
        self.btn_blacklist.clicked.connect(self._show_blacklist_dialog)
        bottom_layout.addWidget(self.btn_blacklist)
        self._update_blacklist_button()

        self.btn_stop = QPushButton("Стоп")
        self.btn_stop.setObjectName("Danger")
        self.btn_stop.clicked.connect(self._on_stop)
        self.btn_stop.setEnabled(False)
        bottom_layout.addWidget(self.btn_stop)

        main_area_layout.addWidget(bottom_bar)

        # Центральный виджет: фильтры сверху, таблица слева, дополнительные фильтры справа.
        page = QWidget()
        cl = QVBoxLayout(page)
        cl.setContentsMargins(0, 0, 0, 0)
        cl.setSpacing(0)
        cl.addWidget(top)
        content = QWidget()
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(10)

        left = QWidget()
        left.setMinimumWidth(0)
        left.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Expanding)
        left_layout = QVBoxLayout(left)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(0)
        left_layout.addWidget(self.sidebar, 0)
        left_layout.addWidget(main_area, 1)
        content_layout.addWidget(left, 1)

        if hasattr(self.sidebar, "extra_scroll"):
            self.sidebar.layout().removeWidget(self.sidebar.extra_scroll)
            self.sidebar.extra_scroll.setParent(content)
            self.sidebar.extra_scroll.setFixedWidth(410)
            content_layout.addWidget(self.sidebar.extra_scroll, 0)

        cl.addWidget(content, 1)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setWidget(page)
        self.setCentralWidget(scroll)

        # Status bar
        sb = QStatusBar()
        self.setStatusBar(sb)
        self.status_msg = QLabel("Готов.")
        sb.addWidget(self.status_msg, 1)
        self.progress = QProgressBar()
        self.progress.setFixedWidth(220)
        self.progress.setRange(0, 0)
        self.progress.setFormat("Обработано: %v из %m")
        self.progress.hide()
        sb.addPermanentWidget(self.progress)

        # Горячие клавиши
        act_search = QAction(self)
        act_search.setShortcut(QKeySequence("Ctrl+Return"))
        act_search.triggered.connect(self._on_search)
        self.addAction(act_search)

        act_focus_query = QAction(self)
        act_focus_query.setShortcut(QKeySequence("Ctrl+F"))
        act_focus_query.triggered.connect(lambda: self.sidebar.ed_query.setFocus())
        self.addAction(act_focus_query)

        self._apply_platform_ui()
        QTimer.singleShot(0, self._apply_table_column_widths)

    def _apply_table_column_widths(self) -> None:
        """Фиксированные ширины колонок — иначе заголовок сжимает их под вьюпорт и скролл не появляется."""
        hh = self.table.horizontalHeader()
        widths = [170, 180, 230, 280, 220, 80, 95, 145, 145, 170, 200]
        n = min(len(widths), self.proxy.columnCount())
        for i in range(n):
            hh.setSectionResizeMode(i, QHeaderView.ResizeMode.Interactive)
            hh.resizeSection(i, widths[i])
        self._schedule_table_row_resize()

    def _build_empty_state(self) -> None:
        self.empty_state = QFrame(self.table)
        self.empty_state.setObjectName("EmptyState")
        self.empty_state.setAttribute(Qt.WidgetAttribute.WA_TransparentForMouseEvents, True)
        box = QVBoxLayout(self.empty_state)
        box.setContentsMargins(0, 0, 0, 0)
        box.setSpacing(8)
        box.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.empty_icon = QLabel()
        self.empty_icon.setAlignment(Qt.AlignmentFlag.AlignCenter)
        pix = QPixmap(str(asset_path("iconfounder.png")))
        if not pix.isNull():
            self.empty_icon.setPixmap(
                pix.scaled(
                    190,
                    140,
                    Qt.AspectRatioMode.KeepAspectRatio,
                    Qt.TransformationMode.SmoothTransformation,
                )
            )
        self.empty_title = QLabel("Результаты поиска появятся здесь")
        self.empty_title.setObjectName("EmptyStateTitle")
        self.empty_title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.empty_text = QLabel("Запустите поиск, чтобы увидеть найденные процедуры")
        self.empty_text.setObjectName("EmptyStateText")
        self.empty_text.setAlignment(Qt.AlignmentFlag.AlignCenter)

        box.addWidget(self.empty_icon)
        box.addWidget(self.empty_title)
        box.addWidget(self.empty_text)
        self.table.horizontalScrollBar().valueChanged.connect(lambda *_: self._position_empty_state())
        self.table.verticalScrollBar().valueChanged.connect(lambda *_: self._position_empty_state())
        self._position_empty_state()
        self._update_empty_state()

    def _position_empty_state(self) -> None:
        if not hasattr(self, "empty_state"):
            return
        margin = 24
        rect = self.table.viewport().geometry()
        self.empty_state.setGeometry(
            rect.x() + margin,
            rect.y() + margin,
            max(0, rect.width() - margin * 2),
            max(0, rect.height() - margin * 2),
        )
        self.empty_state.raise_()

    def _update_empty_state(self) -> None:
        if hasattr(self, "empty_state"):
            self.empty_state.setVisible(self.proxy.rowCount() == 0)

    def _schedule_table_row_resize(self) -> None:
        QTimer.singleShot(0, self._resize_table_rows_to_contents)

    def _resize_table_rows_to_contents(self) -> None:
        self.table.resizeRowsToContents()

    def _clear_api_debug(self) -> None:
        self._api_debug_chunks.clear()
        self.btn_save_api_debug.setEnabled(False)

    @Slot(str)
    def _on_api_debug(self, text: str) -> None:
        if not text:
            return
        self._api_debug_chunks.append(text.strip())
        self.btn_save_api_debug.setEnabled(True)

    def _save_api_debug(self) -> None:
        if not self._api_debug_chunks:
            QMessageBox.information(
                self,
                "Логов пока нет",
                "Запустите поиск. После первых ответов API здесь можно будет сохранить лог.",
            )
            return

        default_name = f"api_logs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Куда сохранить API-логи",
            default_name,
            "JSON (*.json);;Text (*.txt);;Все файлы (*.*)",
        )
        if not path:
            return

        entries: list[Any] = []
        for chunk in self._api_debug_chunks:
            try:
                entries.append(json.loads(chunk))
            except json.JSONDecodeError:
                entries.append({"raw": chunk})
        content = json.dumps(
            {
                "saved_at": datetime.now().isoformat(timespec="seconds"),
                "entries": entries,
            },
            ensure_ascii=False,
            indent=2,
            default=str,
        )
        try:
            Path(path).write_text(content, encoding="utf-8")
        except OSError as e:
            QMessageBox.critical(
                self,
                "Не удалось сохранить лог",
                "Файл не удалось записать. Выберите другую папку или имя файла.\n\n"
                f"Путь: {path}\n\n"
                f"Подробности: {e}",
            )
            return

        QMessageBox.information(self, "API-логи сохранены", f"Файл сохранён:\n{path}")

    def eventFilter(self, watched: QObject, event: QEvent) -> bool:
        if watched is self.table.viewport() and event.type() in {QEvent.Type.Resize, QEvent.Type.Show}:
            self._position_empty_state()
        if watched is self.table.viewport() and event.type() == QEvent.Type.Wheel:
            wheel = event
            if wheel.modifiers() & Qt.KeyboardModifier.ShiftModifier:
                bar = self.table.horizontalScrollBar()
                dy = wheel.angleDelta().y()
                dx = wheel.angleDelta().x()
                step = dx if dx != 0 else dy
                if step != 0:
                    bar.setValue(bar.value() - step)
                    return True
        return super().eventFilter(watched, event)

    # ------------------------------------------------------------------ задачи
    def _is_platform_ready(self) -> bool:
        return self._platform_key in {"gpb", "gpb_business", "roseltorg"}

    def _platform_title(self) -> str:
        if self._platform_key == "roseltorg":
            return "Росэлторг"
        if self._platform_key == "gpb_business":
            return "ГПБ Бизнес"
        return "ЭТП ГПБ"

    def _set_platform_buttons(self) -> None:
        self.btn_platform_gpb.setChecked(self._platform_key in {"gpb", "gpb_business"})
        self.btn_platform_roseltorg.setChecked(self._platform_key == "roseltorg")
        self.btn_platform_gpb.setText(self._platform_title() if self._platform_key in {"gpb", "gpb_business"} else "ГПБ")
        self.act_platform_etp_gpb.setCheckable(True)
        self.act_platform_gpb_business.setCheckable(True)
        self.act_platform_etp_gpb.setChecked(self._platform_key == "gpb")
        self.act_platform_gpb_business.setChecked(self._platform_key == "gpb_business")

    def _apply_platform_ui(self) -> None:
        self._set_platform_buttons()
        if self._platform_key == "roseltorg":
            self.sidebar.set_platform_filter_options(
                ROSELTORG_PROCEDURE_TYPE_OPTIONS,
                ROSELTORG_STATUS_OPTIONS,
                ROSELTORG_SEARCH_BY_OPTIONS,
                platform_key="roseltorg",
            )
            self.title_label.setText("Росэлторг — Актуальные процедуры")
            self.subtitle_label.setText("Поиск, фильтры и ключевые слова")
            self.lbl_counter.setText("Данных нет. Нажмите «Поиск». Если сессии нет, войдите через ЭЦП.")
            self.user_label.setText("Пользователь: —")
            self._set_badge("idle", "○  Росэлторг")
            self.status_msg.setText("Готов. Нажмите «Поиск» и войдите через ЭЦП при необходимости.")
        elif self._platform_key == "gpb_business":
            self.sidebar.set_platform_filter_options(
                PROCEDURE_TYPE_OPTIONS,
                STATUS_OPTIONS,
                None,
                platform_key="gpb_business",
            )
            self.title_label.setText("ГПБ Бизнес — Актуальные процедуры")
            self.subtitle_label.setText("Поиск, фильтры и экспорт")
            if not self.model.rowCount():
                self.lbl_counter.setText("Данных нет. Нажмите «Поиск».")
            self._set_badge("idle", "○  Браузер не запущен")
            self.status_msg.setText("Готов. Нажмите «Поиск».")
        else:
            self.sidebar.set_platform_filter_options(
                PROCEDURE_TYPE_OPTIONS,
                STATUS_OPTIONS,
                None,
                platform_key="gpb",
            )
            self.title_label.setText("ЭТП ГПБ — Актуальные процедуры")
            self.subtitle_label.setText("Поиск, фильтры и экспорт")
            if not self.model.rowCount():
                self.lbl_counter.setText("Данных нет. Нажмите «Поиск».")
            self._set_badge("idle", "○  Браузер не запущен")
            self.status_msg.setText("Готов. Нажмите «Поиск».")

    def _select_platform(self, key: str) -> None:
        if key not in {"gpb", "gpb_business", "roseltorg"}:
            return
        if self.runner.is_running():
            self._set_platform_buttons()
            QMessageBox.information(
                self,
                "Идёт операция",
                "Дождитесь завершения текущей операции перед сменой площадки.",
            )
            return
        if key == self._platform_key:
            self._apply_platform_ui()
            self._update_controls()
            return

        if self._cache_dirty:
            self._save_cache_now()
        self._cache_save_timer.stop()
        self._cache_dirty = False
        try:
            self.client.close()
        except Exception:
            traceback.print_exc()

        self._platform_key = key
        if key == "roseltorg":
            self.client = RoseltorgClient()
        elif key == "gpb_business":
            self.client = GpbBusinessClient()
        else:
            self.client = EtpClient()
        self._set_platform_buttons()
        self.model.clear()
        self._last_total = 0
        self._current_start = 0
        self._last_user = None
        self._refresh_counter()
        self._apply_platform_ui()
        self._update_controls()

    def _ensure_platform_ready(self) -> bool:
        if self._is_platform_ready():
            return True
        QMessageBox.information(
            self,
            "Площадка в разработке",
            f"{self._platform_title()} пока добавлена только как переключатель. "
            "Поиск и загрузка документов будут подключены следующим этапом.",
        )
        return False

    def _apply_selected_browser(self) -> None:
        browser = self.sidebar.selected_browser()
        self.client.configure_browser(
            key=browser.key,
            label=browser.label,
            exe_path=browser.exe_path,
            user_data_dir=browser.user_data_dir,
            profile_dir=browser.profile_dir,
            port=browser.port,
        )

    def _on_search(self) -> None:
        if self.runner.is_running():
            return
        if not self._ensure_platform_ready():
            return

        self._apply_selected_browser()
        filters = self.sidebar.client_filters()
        if filters.keyword_search_enabled and not filters.keywords:
            QMessageBox.information(
                self,
                "Нет ключевых слов",
                "Список ключевых слов пуст. Добавьте слова через «Редактировать список».",
            )
            return
        self.proxy.set_filters(self._display_filters_for_platform(filters))
        self.model.set_keywords(filters.keywords, filters.keyword_lemma_enabled)

        has_active_filters = self._has_active_filters(filters)
        if self._platform_key == "gpb" and CACHE_FILE.exists() and not has_active_filters:
            choice = self._ask_cache_choice()
            if choice == "cancel":
                return
            if choice == "cache":
                self._use_cache()
                return
            # choice == "refresh" → удаляем кэш и идём парсить заново
            self._delete_cache()

        self.model.clear()
        self._current_start = 0
        self._last_total = 0
        self._refresh_counter()
        self._start_task(
            self.sidebar.search_params(),
            start=0,
            batches=self._search_batches(filters),
            filters=filters,
        )

    def _on_edit_keywords(self) -> None:
        dialog = QDialog(self)
        dialog.setWindowTitle("Ключевые слова")
        dialog.resize(720, 560)

        layout = QVBoxLayout(dialog)
        hint = QLabel(
            "Отметьте галочками ключевые слова, по которым нужно искать. "
            "Поиск найдёт процедуры, где встречается хотя бы одно активное слово."
        )
        hint.setWordWrap(True)
        layout.addWidget(hint)

        keyword_list = QListWidget()
        keyword_list.setAlternatingRowColors(True)
        keyword_list.setSelectionMode(QListWidget.ExtendedSelection)
        try:
            keyword_items = load_keyword_items()
        except OSError as e:
            QMessageBox.critical(
                self,
                "Не удалось открыть список",
                "Не удалось открыть файл ключевых слов. "
                "Проверьте доступ к локальной папке приложения.\n\n"
                f"Путь: {KEYWORDS_FILE}\n\n"
                f"Подробности: {e}",
            )
            return
        for enabled, keyword in keyword_items:
            item = QListWidgetItem(keyword)
            item.setFlags(
                item.flags()
                | Qt.ItemIsUserCheckable
                | Qt.ItemIsEditable
                | Qt.ItemIsEnabled
                | Qt.ItemIsSelectable
            )
            item.setCheckState(Qt.Checked if enabled else Qt.Unchecked)
            keyword_list.addItem(item)
        layout.addWidget(keyword_list, 1)

        actions = QHBoxLayout()
        btn_all = QPushButton("Все")
        btn_none = QPushButton("Снять все")
        btn_add = QPushButton("Добавить")
        btn_remove = QPushButton("Удалить выбранные")
        actions.addWidget(btn_all)
        actions.addWidget(btn_none)
        actions.addStretch(1)
        actions.addWidget(btn_add)
        actions.addWidget(btn_remove)
        layout.addLayout(actions)

        def set_all(state: Qt.CheckState) -> None:
            for i in range(keyword_list.count()):
                keyword_list.item(i).setCheckState(state)

        def add_keyword() -> None:
            text, ok = QInputDialog.getText(
                dialog,
                "Добавить ключевое слово",
                "Ключевое слово или фраза:",
            )
            if not ok:
                return
            # parse_keywords() отбрасывает строки с [ ] (выключено) — для «Добавить»
            # нужна любая распознанная фраза; активность задаётся галочкой в списке.
            rows = parse_keyword_items(text.strip())
            if not rows:
                QMessageBox.information(
                    dialog,
                    "Не добавлено",
                    "Текст не принят: пустая строка, служебный фрагмент или слишком короткая фраза "
                    "(до 2 символов, если это не аббревиатура заглавными буквами).",
                )
                return
            _, keyword = rows[0]
            item = QListWidgetItem(keyword)
            item.setFlags(
                item.flags()
                | Qt.ItemIsUserCheckable
                | Qt.ItemIsEditable
                | Qt.ItemIsEnabled
                | Qt.ItemIsSelectable
            )
            item.setCheckState(Qt.Checked)
            keyword_list.addItem(item)

        def remove_selected() -> None:
            for item in keyword_list.selectedItems():
                keyword_list.takeItem(keyword_list.row(item))

        btn_all.clicked.connect(lambda: set_all(Qt.Checked))
        btn_none.clicked.connect(lambda: set_all(Qt.Unchecked))
        btn_add.clicked.connect(add_keyword)
        btn_remove.clicked.connect(remove_selected)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Save
            | QDialogButtonBox.StandardButton.Cancel
        )
        layout.addWidget(buttons)
        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)

        if dialog.exec() != QDialog.Accepted:
            return

        items: list[tuple[bool, str]] = []
        for i in range(keyword_list.count()):
            item = keyword_list.item(i)
            rows = parse_keyword_items(item.text())
            if not rows:
                continue
            _, keyword = rows[0]
            items.append((item.checkState() == Qt.Checked, keyword))
        try:
            save_keyword_items(items)
        except OSError as e:
            QMessageBox.critical(
                self,
                "Не удалось сохранить",
                "Не удалось записать файл ключевых слов (нет прав или диск недоступен).\n\n"
                f"Путь: {KEYWORDS_FILE}\n\n"
                f"Подробности: {e}",
            )
            return
        active_keywords = tuple(keyword for enabled, keyword in items if enabled)
        filters = self.sidebar.client_filters()
        self.model.set_keywords(active_keywords, filters.keyword_lemma_enabled)
        self.sidebar.refresh_keywords_count()
        self._on_filters_changed()
        QMessageBox.information(
            self,
            "Список сохранён",
            f"Активных ключевых слов/фраз: {len(active_keywords)} из {len(items)}.",
        )

    def _ask_cache_choice(self) -> str:
        """Диалог «Показать из кэша / Загрузить заново / Отмена».

        Возвращает одно из: 'cache', 'refresh', 'cancel'.
        """
        meta = self._read_cache_meta()
        box = QMessageBox(self)
        box.setIcon(QMessageBox.Question)
        box.setWindowTitle("Найден сохранённый результат")
        if meta:
            saved_at = (meta.get("saved_at") or "")[:16].replace("T", " ")
            count = meta.get("count") or 0
            box.setText(
                f"В кэше уже есть <b>{count}</b> процедур, сохранённых "
                f"<b>{saved_at}</b>."
            )
        else:
            box.setText("В кэше уже есть сохранённый результат поиска.")
        box.setInformativeText(
            "Что делать?\n\n"
            "• Показать из кэша — мгновенно вывести сохранённые данные.\n"
            "• Загрузить заново — очистить кэш и спарсить с сайта."
        )
        btn_cache = box.addButton("Показать из кэша", QMessageBox.AcceptRole)
        btn_refresh = box.addButton("Загрузить заново", QMessageBox.DestructiveRole)
        btn_cancel = box.addButton("Отмена", QMessageBox.RejectRole)
        box.setDefaultButton(btn_cache)
        box.exec()
        clicked = box.clickedButton()
        if clicked is btn_cache:
            return "cache"
        if clicked is btn_refresh:
            return "refresh"
        return "cancel"

    def _read_cache_meta(self) -> Optional[dict[str, Any]]:
        try:
            data = json.loads(CACHE_FILE.read_text(encoding="utf-8"))
            return {
                "saved_at": data.get("saved_at"),
                "count": len(data.get("procedures") or []),
                "total": data.get("total"),
            }
        except Exception:
            return None

    def _use_cache(self) -> None:
        """Загрузить результат из кэша в таблицу, без обращения к сайту."""
        if not CACHE_FILE.exists():
            return
        try:
            data = json.loads(CACHE_FILE.read_text(encoding="utf-8"))
        except Exception as e:
            QMessageBox.warning(
                self, "Кэш повреждён",
                f"Не удалось прочитать кэш:\n{e}\n\nБудет выполнен новый поиск.",
            )
            self._delete_cache()
            self.model.clear()
            self._current_start = 0
            self._last_total = 0
            self._start_task(
                self.sidebar.search_params(),
                start=0,
                batches=self._search_batches(self.sidebar.client_filters()),
                filters=self.sidebar.client_filters(),
            )
            return
        procs = self._filter_deleted_procedures(data.get("procedures") or [])
        self.model.set_rows(procs)
        self._last_total = int(data.get("total") or len(procs))
        self._current_start = len(procs)
        saved_at = (data.get("saved_at") or "")[:16].replace("T", " ")
        self.status_msg.setText(
            f"Показано из кэша: {len(procs)} процедур (сохранено {saved_at})."
        )
        self._refresh_counter()
        self._schedule_table_row_resize()
        self._update_controls()

    def _delete_cache(self) -> None:
        try:
            if CACHE_FILE.exists():
                CACHE_FILE.unlink()
        except Exception:
            traceback.print_exc()
        self._cache_dirty = False

    def _on_prev_page(self) -> None:
        self.proxy.previous_page()
        self._refresh_counter()
        self.table.scrollToTop()
        self._schedule_table_row_resize()

    def _on_next_page(self) -> None:
        self.proxy.next_page()
        self._refresh_counter()
        self.table.scrollToTop()
        self._schedule_table_row_resize()

    def _selected_procedures(self) -> list[dict[str, Any]]:
        rows = sorted({idx.row() for idx in self.table.selectionModel().selectedRows()})
        if not rows and self.table.currentIndex().isValid():
            rows = [self.table.currentIndex().row()]
        selected: list[dict[str, Any]] = []
        for row in rows:
            src = self.proxy.mapToSource(self.proxy.index(row, 0))
            proc = self.model.row_at(src.row())
            if proc is not None:
                selected.append(proc)
        return selected

    def _on_download_documents(self) -> None:
        if self.runner.is_running():
            return
        if not self._ensure_platform_ready():
            return
        if self._platform_key == "roseltorg":
            QMessageBox.information(
                self,
                "Скачивание документов",
                "Для Росэлторга сейчас реализован поиск процедур. "
                "Скачивание документов подключим отдельным этапом.",
            )
            return
        procedures = self._selected_procedures()
        if not procedures:
            QMessageBox.information(
                self,
                "Ничего не выбрано",
                "Выберите одну или несколько строк в таблице.",
            )
            return
        default_download_dir = DOCUMENTS_DIR
        try:
            default_download_dir.mkdir(parents=True, exist_ok=True)
        except OSError:
            default_download_dir = Path.home() / "Documents"
        output_dir_str = QFileDialog.getExistingDirectory(
            self,
            "Выберите папку для загрузки документов",
            str(default_download_dir),
        )
        if not output_dir_str:
            return
        output_dir = Path(output_dir_str)
        self._apply_selected_browser()

        self.progress.show()
        self.progress.setRange(0, 0)
        self.progress.setFormat("Скачиваю документы...")
        self.btn_stop.setEnabled(True)
        self.sidebar.set_controls_enabled(False)
        self.btn_prev_page.setEnabled(False)
        self.btn_next_page.setEnabled(False)
        self.btn_download_docs.setEnabled(False)
        self.btn_analyze.setEnabled(False)
        self._set_badge("idle", "● Скачиваю документы…")

        fn = make_download_documents_task(self.client, procedures, output_dir)
        try:
            self.runner.start(
                fn,
                on_progress=self._on_progress,
                on_session=self._on_documents_status,
                on_error=self._on_error,
                on_done=self._on_task_done,
            )
        except Exception as e:
            self._on_error(f"Не удалось запустить скачивание: {e}")
            self._on_task_done()

    def _on_analyze_procedures(self) -> None:
        if self.runner.is_running():
            return
        if not self._ensure_platform_ready():
            return
        if self._platform_key == "roseltorg":
            QMessageBox.information(
                self,
                "Анализ",
                "Анализ карточки через LM Studio сейчас доступен только для ЭТП ГПБ.",
            )
            return
        procedures = self._selected_procedures()
        if not procedures:
            QMessageBox.information(
                self,
                "Ничего не выбрано",
                "Выберите одну или несколько строк в таблице.",
            )
            return
        self._apply_selected_browser()
        self._analysis_sink.clear()
        self.btn_show_analysis.setEnabled(False)

        self.progress.show()
        self.progress.setRange(0, 0)
        self.progress.setFormat("Анализирую...")
        self.btn_stop.setEnabled(True)
        self.sidebar.set_controls_enabled(False)
        self.btn_prev_page.setEnabled(False)
        self.btn_next_page.setEnabled(False)
        self.btn_download_docs.setEnabled(False)
        self.btn_analyze.setEnabled(False)
        self.btn_show_analysis.setEnabled(False)
        self._set_badge("idle", "● Анализ карточки и LM Studio…")

        fn = make_analyze_procedure_task(
            self.client,
            procedures,
            LM_STUDIO_BASE_URL,
            LM_STUDIO_MODEL,
            self._analysis_sink,
        )
        try:
            self.runner.start(
                fn,
                on_progress=self._on_progress,
                on_session=self._on_analyze_session,
                on_error=self._on_error,
                on_done=self._on_analyze_task_done,
            )
        except Exception as e:
            self._on_error(f"Не удалось запустить анализ: {e}")
            self._on_task_done()

    @Slot(bool, str)
    def _on_analyze_session(self, ok: bool, message: str) -> None:
        if ok:
            self._set_badge("true", "● Анализ выполнен")
            self.status_msg.setText(message)
        else:
            self._set_badge("false", "⚠ Ошибка анализа")
            self.status_msg.setText(message)

    @Slot()
    def _on_analyze_task_done(self) -> None:
        rows = self._analysis_sink.get("rows") or []
        self._on_task_done()
        if rows:
            try:
                summary_rows = self._save_analysis_tables(rows)
            except Exception as e:
                self._on_error(f"Не удалось сохранить файлы анализа: {e}")
                return
            self.btn_show_analysis.setEnabled(True)
            self._show_analysis_table_dialog(summary_rows)

    def _on_show_analysis_result(self) -> None:
        rows = self._analysis_sink.get("summary_rows") or []
        if not rows:
            QMessageBox.information(self, "Результат анализа", "Пока нет сохранённого результата анализа.")
            return
        self._show_analysis_table_dialog(rows)

    def _safe_analysis_filename(self, name: str, default: str = "analysis") -> str:
        clean = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", name).strip(" .")
        return (clean[:160] or default) + ".docx"

    def _safe_folder_name(self, name: str, default: str = "procedure") -> str:
        clean = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", str(name or "")).strip(" .")
        return clean[:120] or default

    def _save_analysis_tables(self, rows: list[list[str]]) -> list[list[str]]:
        from docx import Document

        ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)
        title_by_registry = self._analysis_sink.get("title_by_registry") or {}
        unpacked_by_registry = self._analysis_sink.get("unpacked_docs_by_registry") or {}
        summary_rows: list[list[str]] = []

        for row in rows:
            registry = str(row[0] if len(row) > 0 else "").strip() or "unknown"
            parsed_title = str(row[4] if len(row) > 4 else "").strip()
            source_title = str(title_by_registry.get(registry) or "").strip()
            title = parsed_title if parsed_title and parsed_title not in {"—", "не указано"} else source_title
            filename = self._safe_analysis_filename(f"{registry}_{title[:80]}", registry)
            path = ANALYSIS_DIR / filename
            n = 2
            while path.exists():
                path = ANALYSIS_DIR / self._safe_analysis_filename(f"{registry}_{title[:70]}_{n}", registry)
                n += 1

            doc = Document()
            doc.add_heading(f"Анализ закупки {registry}", level=1)
            if title:
                doc.add_paragraph(title)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Поле"
            hdr[1].text = "Значение"
            for cell in hdr:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for header, value in zip(ANALYSIS_TABLE_HEADERS_RU, row):
                cells = table.add_row().cells
                cells[0].text = str(header)
                cells[1].text = str(value or "—")

            doc.save(path)
            try:
                from openpyxl import Workbook

                xlsx_path = path.with_suffix(".xlsx")
                wb = Workbook()
                ws = wb.active
                ws.title = "Анализ"
                ws.append(["Реестровый номер", registry])
                if title:
                    ws.append(["Наименование", title])
                ws.append([])
                for header, value in zip(ANALYSIS_TABLE_HEADERS_RU, row):
                    ws.append([str(header), str(value if value is not None else "—")])
                wb.save(xlsx_path)
            except Exception:
                pass

            unpacked_path = Path(str(unpacked_by_registry.get(registry) or ""))
            if not unpacked_path.is_dir():
                fallback_unpacked = ANALYSIS_DIR / "разархивированные_документы" / self._safe_folder_name(registry)
                if fallback_unpacked.is_dir():
                    unpacked_path = fallback_unpacked
            chat_dir = ANALYSIS_DIR / "rag_debug" / self._safe_folder_name(registry)
            summary_rows.append(
                [
                    registry,
                    title or "—",
                    str(path),
                    str(unpacked_path) if unpacked_path else "",
                    str(chat_dir),
                ]
            )

        self._analysis_sink["summary_rows"] = summary_rows
        return summary_rows

    def _show_analysis_table_dialog(self, rows: list[list[str]]) -> None:
        dlg = QDialog(self)
        n = len(rows)
        dlg.setWindowTitle("Результат анализа карточки ЭТП ГПБ" + (f" ({n} процедур)" if n != 1 else ""))
        dlg.resize(min(1100, self.width() + 80), min(520, self.height()))
        layout = QVBoxLayout(dlg)
        hint = QLabel(
            "Полная таблица анализа подготовлена в Word (.docx) и Excel (.xlsx). "
            "Нажмите «Скачать», выберите папку на своём компьютере, и приложение скопирует файлы туда."
        )
        hint.setWordWrap(True)
        layout.addWidget(hint)

        headers = ["Реестровый номер", "Наименование", "Таблица анализа", "Разархивированные документы", "Чат"]
        table = QTableWidget(len(rows), len(headers))
        table.setHorizontalHeaderLabels(headers)
        hh = table.horizontalHeader()
        hh.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        hh.setStretchLastSection(False)
        table.setColumnWidth(0, 150)
        table.setColumnWidth(1, 520)
        table.setColumnWidth(2, 130)
        table.setColumnWidth(3, 180)
        table.setColumnWidth(4, 110)
        table.setWordWrap(True)
        table.setAlternatingRowColors(True)
        table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)

        def centered_cell_widget(widget: QWidget) -> QWidget:
            holder = QWidget()
            holder_lay = QHBoxLayout(holder)
            holder_lay.setContentsMargins(8, 4, 8, 4)
            holder_lay.setSpacing(0)
            holder_lay.addWidget(widget, 0, Qt.AlignmentFlag.AlignCenter)
            return holder

        def row_docs_dir(row: list[str]) -> Path:
            source = str(row[3] if len(row) > 3 else "").strip()
            registry = str(row[0] if row else "").strip()
            candidates = []
            if source:
                candidates.append(Path(source))
            if registry:
                candidates.append(ANALYSIS_DIR / "разархивированные_документы" / self._safe_folder_name(registry))
            for candidate in candidates:
                if candidate.is_dir():
                    return candidate
            return candidates[0] if candidates else Path("__missing_docs_path__")

        display_issues = list(self._analysis_sink.get("document_issues") or [])
        for row in rows:
            registry = str(row[0] if row else "").strip()
            docs_dir = row_docs_dir(row)
            if not docs_dir.is_dir():
                display_issues.append(
                    {
                        "severity": "critical",
                        "registry": registry,
                        "file": "",
                        "message": f"Папка разархивированных документов не найдена: {docs_dir or 'путь не задан'}",
                    }
                )
            elif not any(p.is_file() for p in docs_dir.rglob("*")):
                display_issues.append(
                    {
                        "severity": "critical",
                        "registry": registry,
                        "file": "",
                        "message": f"Папка разархивированных документов пуста: {docs_dir}",
                    }
                )

        def unique_destination(target: Path) -> Path:
            if not target.exists():
                return target
            stem, suffix = target.stem, target.suffix
            n = 2
            while True:
                candidate = target.with_name(f"{stem}_{n}{suffix}")
                if not candidate.exists():
                    return candidate
                n += 1

        def copy_analysis_files(row_index: int) -> None:
            source = rows[row_index][2] if 0 <= row_index < len(rows) and len(rows[row_index]) > 2 else ""
            if not source:
                return
            destination_dir = QFileDialog.getExistingDirectory(
                dlg,
                "Куда скачать таблицу анализа",
                str(Path.home()),
            )
            if not destination_dir:
                return
            try:
                copied: list[Path] = []
                src_docx = Path(source)
                candidates = [src_docx, src_docx.with_suffix(".xlsx")]
                for src in candidates:
                    if not src.is_file():
                        continue
                    dst = unique_destination(Path(destination_dir) / src.name)
                    if src.resolve() != dst.resolve():
                        shutil.copy2(src, dst)
                    copied.append(dst)
                QMessageBox.information(
                    dlg,
                    "Скачивание завершено",
                    "Скопированы файлы:\n" + "\n".join(str(p) for p in copied),
                )
            except Exception as e:
                QMessageBox.critical(dlg, "Ошибка скачивания", f"Не удалось скачать таблицу анализа:\n{e}")

        def copy_unpacked_documents(row_index: int) -> None:
            if not (0 <= row_index < len(rows)):
                return
            src_dir = row_docs_dir(rows[row_index])
            destination_dir = QFileDialog.getExistingDirectory(
                dlg,
                "Куда скачать разархивированные документы",
                str(Path.home()),
            )
            if not destination_dir:
                return
            try:
                if not src_dir.is_dir():
                    raise RuntimeError(f"Папка не найдена: {src_dir}")
                dst_dir = unique_destination(Path(destination_dir) / src_dir.name)
                shutil.copytree(src_dir, dst_dir)
                QMessageBox.information(dlg, "Скачивание завершено", f"Документы скопированы в:\n{dst_dir}")
            except Exception as e:
                QMessageBox.critical(dlg, "Ошибка скачивания", f"Не удалось скачать документы:\n{e}")

        def open_analysis_chat(row_index: int) -> None:
            if not (0 <= row_index < len(rows)):
                return
            row = rows[row_index]
            registry = str(row[0] if len(row) > 0 else "")
            title = str(row[1] if len(row) > 1 else "")
            docs_dir = row_docs_dir(row)
            index_dir = Path(str(row[4] if len(row) > 4 else ""))
            has_index = (index_dir / "index.faiss").is_file() and (index_dir / "metadata.json").is_file()
            if not has_index and not docs_dir.is_dir():
                QMessageBox.information(
                    dlg,
                    "Чат недоступен",
                    "FAISS-индекс для этой процедуры не найден. "
                    "Также не найдена папка документов, из которой можно построить индекс заново.",
                )
                return

            chat_dlg = QDialog(dlg)
            chat_dlg.setWindowTitle(f"Чат по анализу {registry}")
            chat_dlg.resize(820, 620)
            chat_layout = QVBoxLayout(chat_dlg)
            header = QLabel(f"<b>{registry}</b><br>{title}")
            header.setWordWrap(True)
            chat_layout.addWidget(header)
            history = QTextEdit()
            history.setReadOnly(True)
            history.setPlainText(
                "Задайте вопрос по карточке и документам закупки. "
                "Будут отправлены только самые релевантные фрагменты из FAISS.\n"
            )
            chat_layout.addWidget(history, 1)
            question = QTextEdit()
            question.setPlaceholderText("Например: какие сроки поставки и условия оплаты?")
            question.setMaximumHeight(86)
            chat_layout.addWidget(question)
            buttons_row = QHBoxLayout()
            buttons_row.addStretch(1)
            btn_ask = QPushButton("Спросить")
            btn_close = QPushButton("Закрыть")
            buttons_row.addWidget(btn_ask)
            buttons_row.addWidget(btn_close)
            chat_layout.addLayout(buttons_row)

            def ask() -> None:
                q = question.toPlainText().strip()
                if not q:
                    return
                question.clear()
                history.append(f"\nПользователь:\n{q}\n")
                btn_ask.setEnabled(False)
                QApplication.setOverrideCursor(Qt.WaitCursor)
                try:
                    answer = answer_question_from_saved_index(
                        index_dir=index_dir,
                        question=q,
                        lm_base_url=LM_STUDIO_BASE_URL,
                        lm_model=LM_STUDIO_MODEL,
                        fallback_docs_dir=docs_dir,
                    )
                    history.append(f"Бот:\n{answer}\n")
                except Exception as e:
                    history.append(f"Ошибка:\n{e}\n")
                finally:
                    QApplication.restoreOverrideCursor()
                    btn_ask.setEnabled(True)

            btn_ask.clicked.connect(ask)
            btn_close.clicked.connect(chat_dlg.reject)
            chat_dlg.exec()

        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                item = QTableWidgetItem(val)
                item.setToolTip(val[:2000] if val else "")
                if c in {2, 3, 4}:
                    item.setText("")
                if c == 1:
                    item.setTextAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
                table.setItem(r, c, item)

            btn_table = QPushButton("Скачать")
            btn_table.setEnabled(bool(row[2] if len(row) > 2 else ""))
            btn_table.clicked.connect(lambda _checked=False, row_index=r: copy_analysis_files(row_index))
            table.setCellWidget(r, 2, centered_cell_widget(btn_table))

            btn_docs = QPushButton("Скачать")
            docs_dir = row_docs_dir(row)
            btn_docs.setEnabled(docs_dir.is_dir() and any(p.is_file() for p in docs_dir.rglob("*")))
            if not btn_docs.isEnabled():
                btn_docs.setToolTip(f"Документы недоступны: {docs_dir or 'путь не задан'}")
            btn_docs.clicked.connect(lambda _checked=False, row_index=r: copy_unpacked_documents(row_index))
            table.setCellWidget(r, 3, centered_cell_widget(btn_docs))

            btn_chat = QPushButton("Чат")
            chat_source = Path(str(row[4] if len(row) > 4 else ""))
            docs_source = row_docs_dir(row)
            btn_chat.setEnabled(
                (
                    (chat_source / "index.faiss").is_file()
                    and (chat_source / "metadata.json").is_file()
                )
                or docs_source.is_dir()
            )
            btn_chat.clicked.connect(lambda _checked=False, row_index=r: open_analysis_chat(row_index))
            table.setCellWidget(r, 4, centered_cell_widget(btn_chat))

            title_len = len(str(row[1] if len(row) > 1 else ""))
            table.setRowHeight(r, min(112, max(72, 54 + (title_len // 90) * 18)))

        layout.addWidget(table, 1)

        issues_label = QLabel("Ошибки обработки документов")
        issues_label.setStyleSheet("font-weight: 600; margin-top: 6px;")
        layout.addWidget(issues_label)

        if display_issues:
            issue_table = QTableWidget(len(display_issues), 4)
            issue_table.setHorizontalHeaderLabels(["!", "Важность", "Реестровый номер", "Описание"])
            issue_table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
            issue_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
            issue_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Fixed)
            issue_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
            issue_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
            issue_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
            issue_table.setColumnWidth(0, 28)
            issue_table.setMaximumHeight(150)
            for r, issue in enumerate(display_issues):
                severity = str(issue.get("severity") or "important")
                is_critical = severity == "critical"
                color = QColor("#c00000" if is_critical else "#b8860b")
                level_text = "Критичная" if is_critical else "Важная"
                file_text = str(issue.get("file") or "").strip()
                message = str(issue.get("message") or "").strip()
                if file_text:
                    message = f"{file_text}: {message}"
                values = ["!", level_text, str(issue.get("registry") or ""), message]
                for c, val in enumerate(values):
                    item = QTableWidgetItem(val)
                    item.setToolTip(val)
                    if c in {0, 1}:
                        font = QFont(item.font())
                        font.setBold(True)
                        item.setFont(font)
                        item.setForeground(color)
                    issue_table.setItem(r, c, item)
            layout.addWidget(issue_table)
        else:
            no_issues = QLabel("Ошибок обработки документов нет.")
            no_issues.setStyleSheet("color: #4a515a;")
            layout.addWidget(no_issues)

        raw_map = self._analysis_sink.get("raw_by_registry") or {}
        if raw_map:

            def show_raw() -> None:
                raw_dlg = QDialog(dlg)
                raw_dlg.setWindowTitle("Сырой ответ модели")
                raw_dlg.resize(900, 600)
                rl = QVBoxLayout(raw_dlg)
                te = QTextEdit()
                te.setReadOnly(True)
                te.setPlainText(json.dumps(raw_map, ensure_ascii=False, indent=2))
                rl.addWidget(te)
                bb = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
                bb.rejected.connect(raw_dlg.reject)
                rl.addWidget(bb)
                raw_dlg.exec()

            btn_raw = QPushButton("Сырой ответ модели…")
            btn_raw.clicked.connect(show_raw)
            layout.addWidget(btn_raw)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        buttons.rejected.connect(dlg.reject)
        layout.addWidget(buttons)
        dlg.exec()

    def _search_batches(self, filters: ClientFilters) -> int:
        return 10_000

    def _has_active_filters(self, filters: ClientFilters) -> bool:
        return any(
            (
                bool(filters.quick_search),
                bool(filters.keyword_search_enabled),
                bool(filters.registry_contains),
                bool(filters.unique_number_contains),
                bool(filters.organizer_contains),
                bool(filters.customer_contains),
                bool(filters.customer_region_contains),
                bool(filters.customer_agent_contains),
                bool(filters.title_contains),
                bool(filters.okpd2_contains),
                bool(filters.okved2_contains),
                filters.guarantee_min is not None,
                filters.guarantee_max is not None,
                bool(filters.responsible_contains),
                bool(filters.trend_pur),
                bool(filters.step_ids),
                bool(filters.purchase_form),
                filters.applics_min is not None,
                filters.applics_max is not None,
                filters.lots_min is not None,
                filters.lots_max is not None,
                filters.price_min is not None,
                filters.price_max is not None,
                filters.published_from is not None,
                filters.published_to is not None,
                filters.end_from is not None,
                filters.end_to is not None,
                filters.results_from is not None,
                filters.results_to is not None,
                bool(filters.special_features_contains),
                bool(filters.position_name_contains),
                bool(filters.national_regime_contains),
            )
        )

    def _start_task(
        self,
        params: SearchParams,
        start: int,
        batches: int,
        filters: Optional[ClientFilters] = None,
    ) -> None:
        filters = filters if filters is not None else self.sidebar.client_filters()
        self._cache_save_enabled = not self._has_active_filters(filters)
        self.progress.show()
        self.progress.setRange(0, 0)
        self.progress.setFormat("Ищу процедуры...")
        self._clear_api_debug()
        self.btn_stop.setEnabled(True)
        self.sidebar.set_controls_enabled(False)
        self.btn_prev_page.setEnabled(False)
        self.btn_next_page.setEnabled(False)
        self.btn_download_docs.setEnabled(False)
        self.btn_show_analysis.setEnabled(False)
        self.btn_analyze.setEnabled(False)
        self._set_badge("idle", "● Работаю…")

        fn = make_search_task(
            self.client,
            params,
            start,
            batches,
            client_filters=filters,
        )
        try:
            self.runner.start(
                fn,
                on_progress=self._on_progress,
                on_session=self._on_session_status,
                on_batch=self._on_batch_loaded,
                on_debug=self._on_api_debug,
                on_error=self._on_error,
                on_done=self._on_task_done,
            )
        except Exception as e:
            self._on_error(f"Не удалось запустить задачу: {e}")
            self._on_task_done()

    # --------------- слоты от Worker
    @Slot(str)
    def _on_progress(self, text: str) -> None:
        self.status_msg.setText(text)

    @Slot(bool, str)
    def _on_session_status(self, ok: bool, message: str) -> None:
        if ok:
            self._set_badge("true", "● Сессия активна")
            # Подхватим логин пользователя
            try:
                login = self.client.current_user_login()
            except Exception:
                login = None
            if login:
                self._last_user = login
                self.user_label.setText(f"Пользователь: {login}")
        else:
            self._set_badge("false", "○ Нужен вход")
            # Диалог с подсказкой + кнопкой «Повторить»
            box = QMessageBox(self)
            box.setIcon(QMessageBox.Information)
            box.setWindowTitle("Требуется авторизация")
            box.setText("Сессия не активна.")
            box.setInformativeText(message)
            btn_retry = box.addButton("Я вошёл — повторить", QMessageBox.AcceptRole)
            box.addButton("Отмена", QMessageBox.RejectRole)
            box.exec()
            if box.clickedButton() is btn_retry:
                QTimer.singleShot(200, self._on_search)

    @Slot(bool, str)
    def _on_documents_status(self, ok: bool, message: str) -> None:
        self._set_badge("true" if ok else "false", "● Документы скачаны" if ok else "⚠ Ошибка")
        self.status_msg.setText(message)
        QMessageBox.information(self, "Скачивание документов", message)

    @Slot(list, int, int)
    def _on_batch_loaded(self, procs: list, start: int, total: int) -> None:
        self._last_total = total or self._last_total
        procs = self._filter_deleted_procedures(procs)
        if total:
            processed = min(max(start, 0), total)
            self.progress.setRange(0, total)
            self.progress.setValue(processed)
            self.progress.setFormat(f"Обработано: {processed} из {total}")
        if start == 0 and self.model.rowCount() == 0:
            self.model.set_rows(procs)
        else:
            self.model.append_rows(procs)
        self._current_start = start
        self.proxy.refresh_page()
        self._refresh_counter()
        self._schedule_table_row_resize()
        if procs and self._cache_save_enabled:
            self._schedule_cache_save()

    @Slot(str)
    def _on_error(self, msg: str) -> None:
        # Не даём приложению упасть — только сообщаем пользователю.
        self._set_badge("false", "⚠ Ошибка")
        self.status_msg.setText(msg.splitlines()[0] if msg else "Ошибка")
        box = QMessageBox(self)
        box.setIcon(QMessageBox.Warning)
        box.setWindowTitle("Ошибка")
        box.setText("При выполнении операции возникла ошибка:")
        box.setDetailedText(msg)
        box.exec()

    @Slot()
    def _on_task_done(self) -> None:
        self.progress.hide()
        self.btn_stop.setEnabled(False)
        self.sidebar.set_controls_enabled(True)
        self._update_controls()
        self.status_msg.setText(f"Найдено {self.proxy.filtered_count()} процедур.")

    def _on_stop(self) -> None:
        self.runner.request_stop()
        self.btn_stop.setEnabled(False)
        self.status_msg.setText("Останавливаю…")
        try:
            self.client.close()
        except Exception:
            traceback.print_exc()

    # --------------- клиентские фильтры
    def _display_filters_for_platform(self, filters: ClientFilters) -> ClientFilters:
        if self._platform_key != "roseltorg":
            return filters
        # Росэлторг фильтруется серверным API. В таблице оставляем только
        # ключевые слова, иначе локальный GPB-фильтр скрывает уже найденные строки.
        return replace(
            filters,
            quick_search="",
            registry_contains="",
            unique_number_contains="",
            organizer_contains="",
            customer_contains="",
            customer_region_contains="",
            customer_agent_contains="",
            title_contains="",
            okpd2_contains="",
            okved2_contains="",
            guarantee_min=None,
            guarantee_max=None,
            responsible_contains="",
            trend_pur="",
            trend_pur_values=(),
            step_ids=(),
            law_values=(),
            purchase_form="",
            applics_min=None,
            applics_max=None,
            lots_min=None,
            lots_max=None,
            price_min=None,
            price_max=None,
            published_from=None,
            published_to=None,
            end_from=None,
            end_to=None,
            results_from=None,
            results_to=None,
            special_features_contains="",
            position_name_contains="",
            national_regime_contains="",
        )

    def _on_filters_changed(self) -> None:
        filters = self.sidebar.client_filters()
        self.proxy.set_filters(self._display_filters_for_platform(filters))
        self.model.set_keywords(filters.keywords, filters.keyword_lemma_enabled)
        self._refresh_counter()
        self._schedule_table_row_resize()

    def _on_reset_filters(self) -> None:
        self.sidebar.reset_client_filters()
        self.proxy.set_filters(ClientFilters())
        self.model.set_keywords((), False)
        self._refresh_counter()
        self._schedule_table_row_resize()

    # --------------- таблица
    def _proc_from_index(self, idx: QModelIndex) -> Optional[dict[str, Any]]:
        if not idx.isValid():
            return None
        src = self.proxy.mapToSource(idx)
        return self.model.row_at(src.row())

    def _on_row_double_clicked(self, idx: QModelIndex) -> None:
        self._open_in_browser(self._proc_from_index(idx))

    def _on_context_menu(self, pos) -> None:
        idx = self.table.indexAt(pos)
        if not idx.isValid():
            return
        proc = self._proc_from_index(idx)
        if not proc:
            return
        cell_value = str(idx.data(Qt.DisplayRole) or "").strip()
        menu = QMenu(self)
        copy_cell_action = menu.addAction(
            "Копировать значение ячейки",
            lambda value=cell_value: QApplication.clipboard().setText(value),
        )
        copy_cell_action.setEnabled(bool(cell_value))
        menu.addSeparator()
        menu.addAction("Открыть в Chrome", lambda: self._open_in_browser(proc))
        menu.addSeparator()
        menu.addAction(
            "Копировать реестровый №",
            lambda: QApplication.clipboard().setText(
                str((proc or {}).get("registry_number") or (proc or {}).get("procedure_number") or "")
            ),
        )
        menu.addAction(
            "Копировать наименование",
            lambda: QApplication.clipboard().setText(str((proc or {}).get("title") or "")),
        )
        menu.addAction(
            "Копировать ИНН организатора",
            lambda: QApplication.clipboard().setText(str((proc or {}).get("org_inn") or "")),
        )
        menu.addSeparator()
        delete_action = menu.addAction("Удалить тендер")
        delete_action.triggered.connect(lambda checked=False, p=proc: self._delete_tender(p))
        menu.exec(self.table.viewport().mapToGlobal(pos))

    def _procedure_delete_key(self, proc: dict[str, Any]) -> str:
        for key in ("id", "procedure_id", "registry_number", "procedure_number", "procedure_number2"):
            value = str(proc.get(key) or "").strip()
            if value:
                return f"{self._platform_key}:{key}:{value}"
        title = str(proc.get("title") or proc.get("name") or "").strip()
        organizer = str(proc.get("organizer") or proc.get("org_name") or "").strip()
        return f"{self._platform_key}:fallback:{title}|{organizer}"

    def _deleted_record_title(self, record: dict[str, Any]) -> str:
        proc = record.get("procedure")
        if not isinstance(proc, dict):
            proc = {}
        registry = str(record.get("registry") or proc.get("registry_number") or proc.get("procedure_number") or "").strip()
        title = str(record.get("title") or proc.get("title") or proc.get("name") or "").strip()
        if registry and title:
            return f"{registry} — {title}"
        return registry or title or str(record.get("key") or "Удалённый тендер")

    def _deleted_record_for_proc(self, key: str, proc: dict[str, Any]) -> dict[str, Any]:
        return {
            "key": key,
            "platform": self._platform_key,
            "deleted_at": datetime.now().isoformat(timespec="seconds"),
            "registry": str(proc.get("registry_number") or proc.get("procedure_number") or ""),
            "title": str(proc.get("title") or proc.get("name") or ""),
            "organizer": str(proc.get("organizer") or proc.get("short_name") or proc.get("full_name") or ""),
            "procedure": proc,
        }

    def _load_deleted_tenders(self) -> dict[str, dict[str, Any]]:
        try:
            if not DELETED_TENDERS_FILE.exists():
                return {}
            data = json.loads(DELETED_TENDERS_FILE.read_text(encoding="utf-8"))
            records: dict[str, dict[str, Any]] = {}
            if isinstance(data, dict):
                values = data.get("items") or data.get("deleted") or []
            else:
                values = data
            for value in values:
                if isinstance(value, dict):
                    key = str(value.get("key") or "").strip()
                    if key:
                        records[key] = value
                else:
                    key = str(value).strip()
                    if key:
                        records[key] = {"key": key, "title": key, "procedure": {}}
            return records
        except Exception:
            traceback.print_exc()
            return {}

    def _save_deleted_tenders(self) -> bool:
        try:
            payload = {
                "saved_at": datetime.now().isoformat(timespec="seconds"),
                "items": [
                    self._deleted_tender_records[key]
                    for key in sorted(self._deleted_tender_records)
                ],
            }
            DELETED_TENDERS_FILE.parent.mkdir(parents=True, exist_ok=True)
            DELETED_TENDERS_FILE.write_text(
                json.dumps(payload, ensure_ascii=False, indent=2, default=str),
                encoding="utf-8",
            )
            self._deleted_tender_keys = set(self._deleted_tender_records)
            self._update_blacklist_button()
            return True
        except Exception as e:
            QMessageBox.warning(
                self,
                "Не удалось сохранить удаление",
                f"Список удалённых тендеров не удалось записать на диск.\n\n"
                f"Путь: {DELETED_TENDERS_FILE}\n\n"
                f"Подробности: {e}",
            )
            return False

    def _filter_deleted_procedures(self, procs: list[dict[str, Any]]) -> list[dict[str, Any]]:
        if not self._deleted_tender_keys:
            return procs
        return [proc for proc in procs if self._procedure_delete_key(proc) not in self._deleted_tender_keys]

    def _delete_tender(self, proc: dict[str, Any]) -> None:
        key = self._procedure_delete_key(proc)
        title = str(proc.get("title") or proc.get("name") or proc.get("registry_number") or "этот тендер")
        answer = QMessageBox.question(
            self,
            "Удалить тендер",
            f"Удалить тендер из таблицы и скрывать его при следующих поисках?\n\n{title}",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No,
        )
        if answer != QMessageBox.StandardButton.Yes:
            return
        self._deleted_tender_records[key] = self._deleted_record_for_proc(key, proc)
        self._deleted_tender_keys = set(self._deleted_tender_records)
        if not self._save_deleted_tenders():
            self._deleted_tender_records.pop(key, None)
            self._deleted_tender_keys = set(self._deleted_tender_records)
            return
        self.model.set_rows(self._filter_deleted_procedures(self.model.rows()))
        self.proxy.refresh_page()
        self._refresh_counter()
        self._schedule_table_row_resize()
        self._schedule_cache_save()
        self.status_msg.setText(f"Тендер удалён. Список удалённых сохранён: {DELETED_TENDERS_FILE}")

    def _update_blacklist_button(self) -> None:
        if hasattr(self, "btn_blacklist"):
            count = len(self._deleted_tender_records)
            self.btn_blacklist.setText(f"Черный список ({count})" if count else "Черный список")

    def _show_blacklist_dialog(self) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle("Черный список")
        dlg.resize(760, 480)
        layout = QVBoxLayout(dlg)

        info = QLabel(f"Удалённые тендеры сохраняются здесь:\n{DELETED_TENDERS_FILE}")
        info.setWordWrap(True)
        layout.addWidget(info)

        deleted_list = QListWidget()
        deleted_list.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        for key, record in sorted(
            self._deleted_tender_records.items(),
            key=lambda item: self._deleted_record_title(item[1]).casefold(),
        ):
            item = QListWidgetItem(self._deleted_record_title(record))
            item.setData(Qt.UserRole, key)
            item.setToolTip(key)
            deleted_list.addItem(item)
        layout.addWidget(deleted_list, 1)

        buttons = QDialogButtonBox()
        btn_restore = buttons.addButton("Восстановить выбранные", QDialogButtonBox.ButtonRole.ActionRole)
        btn_close = buttons.addButton("Закрыть", QDialogButtonBox.ButtonRole.RejectRole)
        layout.addWidget(buttons)

        def restore_selected() -> None:
            selected = deleted_list.selectedItems()
            if not selected:
                QMessageBox.information(dlg, "Ничего не выбрано", "Выберите одну или несколько строк для восстановления.")
                return
            restored: list[dict[str, Any]] = []
            for item in selected:
                key = str(item.data(Qt.UserRole) or "")
                record = self._deleted_tender_records.pop(key, None)
                proc = record.get("procedure") if isinstance(record, dict) else None
                if isinstance(proc, dict) and proc:
                    restored.append(proc)
                deleted_list.takeItem(deleted_list.row(item))
            self._deleted_tender_keys = set(self._deleted_tender_records)
            if not self._save_deleted_tenders():
                return
            if restored:
                existing_keys = {self._procedure_delete_key(row) for row in self.model.rows()}
                to_add = [proc for proc in restored if self._procedure_delete_key(proc) not in existing_keys]
                self.model.append_rows(to_add)
            self.proxy.refresh_page()
            self._refresh_counter()
            self._schedule_table_row_resize()
            self._schedule_cache_save()
            self.status_msg.setText(f"Восстановлено: {len(selected)}.")

        btn_restore.clicked.connect(restore_selected)
        btn_close.clicked.connect(dlg.reject)
        dlg.exec()
        self._update_blacklist_button()

    def _open_in_browser(self, proc: Optional[dict[str, Any]]) -> None:
        if not proc:
            return
        pid = proc.get("id")
        if not pid:
            return
        if proc.get("url"):
            webbrowser.open(str(proc["url"]))
            return
        webbrowser.open(VIEW_URL.format(pid=pid))

    # --------------- экспорт
    def _on_export(self) -> None:
        if self.model.rowCount() == 0:
            QMessageBox.information(self, "Нет данных", "Сначала выполните поиск.")
            return
        default_name = f"procedures_{datetime.now():%Y%m%d_%H%M}.xlsx"
        path_str, _ = QFileDialog.getSaveFileName(
            self, "Сохранить в XLSX",
            str(Path.cwd() / default_name),
            "Excel (*.xlsx)",
        )
        if not path_str:
            return
        visible: list[dict[str, Any]] = []
        for i in range(self.proxy.rowCount()):
            src = self.proxy.mapToSource(self.proxy.index(i, 0))
            r = self.model.row_at(src.row())
            if r is not None:
                visible.append(r)
        try:
            self._write_xlsx(Path(path_str), visible)
            QMessageBox.information(
                self, "Готово",
                f"Сохранено {len(visible)} строк в\n{path_str}",
            )
        except Exception as e:
            QMessageBox.critical(self, "Ошибка экспорта", str(e))

    def _write_xlsx(self, path: Path, procs: list[dict[str, Any]]) -> None:
        from openpyxl import Workbook
        from openpyxl.styles import Font

        wb = Workbook()
        ws = wb.active
        ws.title = "Процедуры"
        titles = [c[1] for c in COLUMNS] + ["id", "Дата публикации"]
        ws.append(titles)
        for cell in ws[1]:
            cell.font = Font(bold=True)
        for p in procs:
            ws.append(
                [
                    p.get("registry_number") or p.get("procedure_number") or "",
                    self.model._display(p, "trend_pur_label"),
                    p.get("short_name") or p.get("full_name") or "",
                    p.get("title") or "",
                    ", ".join(self.model._keyword_matches(p)),
                    ", ".join(str(t) for t in (p.get("tags") or [])),
                    p.get("applics_count") or 0,
                    fmt_date(
                        self.model._first_date(
                            p,
                            (
                                "date_start_registration",
                                "date_begin_registration",
                                "date_registration_start",
                                "date_start_applic",
                                "date_begin_applic",
                                "date_published",
                            ),
                        )
                    ),
                    fmt_date(parse_dt(p.get("date_end_registration"))),
                    parse_price(p.get("total_price")) or "",
                    self.model._display(p, "step_label"),
                    p.get("id"),
                    fmt_date(parse_dt(p.get("date_published"))),
                ]
            )
        widths = [18, 28, 28, 80, 30, 12, 12, 20, 20, 18, 28, 10, 20]
        for i, w in enumerate(widths, start=1):
            col_letter = ws.cell(row=1, column=i).column_letter
            ws.column_dimensions[col_letter].width = w
        ws.freeze_panes = "A2"
        wb.save(path)

    # --------------- счётчики / бейдж
    def _refresh_counter(self) -> None:
        loaded = self.model.rowCount()
        visible = self.proxy.rowCount()
        found = self.proxy.filtered_count()
        page_count = self.proxy.page_count()
        current_page = self.proxy.current_page() + 1 if found else 0
        if loaded == 0:
            self.lbl_counter.setText("Данных нет. Нажмите «Поиск».")
            self.lbl_page.setText("Страница 0 из 0")
        elif found == 0:
            self.lbl_counter.setText("По текущим фильтрам ничего не найдено.")
            self.lbl_page.setText("Страница 0 из 0")
        else:
            self.lbl_counter.setText(
                f"Найдено {found}. Показано {visible} на странице."
            )
            self.lbl_page.setText(f"Страница {current_page} из {page_count}")
        if loaded > 0 and hasattr(self, "cache_banner"):
            self.cache_banner.setVisible(False)
        self._update_empty_state()
        self._update_controls()

    def _set_badge(self, state: str, text: str) -> None:
        self.session_badge.setProperty("ok", state)
        self.session_badge.setText(text)
        self.session_badge.style().unpolish(self.session_badge)
        self.session_badge.style().polish(self.session_badge)

    def _update_controls(self) -> None:
        running = self.runner.is_running()
        platform_ready = self._is_platform_ready()
        found = self.proxy.filtered_count()
        has_rows = self.model.rowCount() > 0
        has_visible_rows = found > 0
        current_page = self.proxy.current_page()
        page_count = self.proxy.page_count()
        self.btn_platform_gpb.setEnabled(not running)
        self.btn_platform_roseltorg.setEnabled(not running)
        self.btn_prev_page.setEnabled(platform_ready and current_page > 0)
        self.btn_next_page.setEnabled(platform_ready and current_page + 1 < page_count)
        gpb_like = self._platform_key in {"gpb", "gpb_business"}
        self.btn_download_docs.setEnabled(platform_ready and gpb_like and not running and has_visible_rows)
        self.btn_analyze.setEnabled(platform_ready and gpb_like and not running and has_visible_rows)
        self.btn_show_analysis.setEnabled(not running and bool(self._analysis_sink.get("summary_rows")))
        self.btn_export.setEnabled(platform_ready and has_rows)
        self.sidebar.set_controls_enabled(platform_ready and not running)

    # --------------- кэш
    def _schedule_cache_save(self) -> None:
        self._cache_dirty = True
        self._cache_save_timer.start(1000)

    def _save_cache_now(self) -> None:
        if not self._cache_dirty:
            return
        self._cache_dirty = False
        try:
            CACHE_FILE.parent.mkdir(parents=True, exist_ok=True)
            payload = {
                "saved_at": datetime.now().isoformat(),
                "date_from": self.sidebar.ed_date_from.date().toString("dd.MM.yyyy"),
                "query": self.sidebar.ed_query.text().strip(),
                "total": self._last_total,
                "procedures": self.model.rows(),
            }
            CACHE_FILE.write_text(
                json.dumps(payload, ensure_ascii=False), encoding="utf-8"
            )
        except Exception:
            traceback.print_exc()

    def _announce_cache_on_start(self) -> None:
        """При старте только сообщаем, что есть кэш — не загружаем его автоматически.

        Сам выбор (использовать / перезагрузить) делается при клике «Поиск».
        """
        meta = self._read_cache_meta()
        if meta and meta.get("count"):
            saved_at = (meta.get("saved_at") or "")[:16].replace("T", " ")
            self.lbl_counter.setText(
                f"Есть сохранённый результат: {meta['count']} процедур от {saved_at}. "
                "Нажмите «Поиск», чтобы выбрать действие."
            )
            self.cache_banner.setVisible(True)
            self.status_msg.setText(
                "Готов. Найден кэш — выбор предложат при нажатии «Поиск»."
            )
        else:
            self.cache_banner.setVisible(False)
            self.status_msg.setText("Готов. Нажмите «Поиск».")

    def _cleanup_analysis_temp_dirs(self) -> None:
        """Удаляет только служебные RAG/analysis-папки, сохраняя готовые .docx/.xlsx."""
        temp_dir_names = (
            "rag_debug",
            "_downloaded_docs",
            "разархивированные_документы",
        )
        for name in temp_dir_names:
            path = ANALYSIS_DIR / name
            try:
                if path.is_dir():
                    shutil.rmtree(path, ignore_errors=True)
            except Exception:
                traceback.print_exc()

    # --------------- закрытие
    def closeEvent(self, event) -> None:  # noqa: N802
        try:
            self._save_cache_now()
        except Exception:
            pass
        try:
            self.runner.shutdown(wait_ms=2000)
        except Exception:
            pass
        try:
            self.client.close()
        except Exception:
            pass
        try:
            self._cleanup_analysis_temp_dirs()
        except Exception:
            pass
        event.accept()
